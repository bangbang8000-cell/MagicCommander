from pandas import read_excel
from numpy import int64 as numpy_int64
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENTATION
import os
import logging
from config import WORKSPACE_DIR

logger = logging.getLogger(__name__)


# 打印配置中文字段映射（前端 → Python docx）
_FORMAT_MAP = {
    'A4': Cm(29.7),    # 纵向时作为长度，自动转换
    'A5': Cm(21.0),
}

_ORIENTATION_MAP = {
    'portrait': WD_ORIENTATION.PORTRAIT,
    'landscape': WD_ORIENTATION.LANDSCAPE,
}


def _apply_page_config(doc, config):
    """把 config 应用到 Word 文档的最后一个 section"""
    if not config or not isinstance(config, dict):
        return

    section = doc.sections[-1]

    # 纸张方向
    orientation = config.get('orientation', 'portrait')
    section.orientation = _ORIENTATION_MAP.get(orientation, WD_ORIENTATION.PORTRAIT)

    # 纸张大小（简化：A4 / A5 按 cm 设置 page_width / page_height）
    fmt = config.get('format', 'A4')
    # A4 纵向: 21cm × 29.7cm / A5: 14.8 × 21
    paper_sizes = {
        'A4': (Cm(21.0), Cm(29.7)),
        'A5': (Cm(14.8), Cm(21.0)),
    }
    if fmt in paper_sizes:
        width, height = paper_sizes[fmt]
        if orientation == 'landscape':
            section.page_width = height
            section.page_height = width
        else:
            section.page_width = width
            section.page_height = height

    # 边距（单位：mm → 转 cm）
    margins_cfg = config.get('margins')
    if isinstance(margins_cfg, dict):
        top = margins_cfg.get('top')
        bottom = margins_cfg.get('bottom')
        left = margins_cfg.get('left')
        right = margins_cfg.get('right')
        if top is not None:
            section.top_margin = Cm(float(top) / 10.0)
        if bottom is not None:
            section.bottom_margin = Cm(float(bottom) / 10.0)
        if left is not None:
            section.left_margin = Cm(float(left) / 10.0)
        if right is not None:
            section.right_margin = Cm(float(right) / 10.0)


def exceltolabel(target_list: list, time_str: str, config=None):
    """把项目的主机表 Excel 转换为 Word 标签文档

    :param target_list: 项目名称列表
    :param time_str: 时间字符串（用于命名输出文件）
    :param config: 可选 dict，包含 format/orientation/labelsPerPage/labelSize/margins
    """
    mid_path = WORKSPACE_DIR
    sheet_name = '主机表'
    excel_name = 'hostname.xlsx'
    word_name = time_str + '_label.docx'

    # 从 config 中读取每页标签数量，默认 5 个后分页
    labels_per_page = 5
    if isinstance(config, dict):
        labels_per_page = int(config.get('labelsPerPage', 5))

    for target in target_list:
        output_dir = os.path.join(mid_path, target, 'output-label')
        os.makedirs(output_dir, exist_ok=True)

        excel_path = os.path.join(mid_path, target, 'excel', excel_name)
        if not os.path.exists(excel_path):
            logger.warning(f'{target} 的 excel 中没有 hostname.xlsx 文件，本次不处理该项目')
            continue

        word_path = os.path.join(output_dir, word_name)
        data = read_excel(excel_path, sheet_name=None, keep_default_na=False)

        ans_dict = {}
        header = ['设备名', 'SN', '型号', '角色', '楼层', '机柜', 'U数', '管理IP', '管理接口']

        for sheet1, value in data.items():
            if sheet1 != sheet_name:
                continue
            for i in value.index.values:
                tmp_dict = {}
                for h in header:
                    tmp_dict[h] = str(value[h][i]).strip()
                ans_dict[str(value[header[0]][i]).strip()] = tmp_dict

        d = Document()

        # 应用纸张 / 方向 / 边距
        _apply_page_config(d, config)

        sum_label = 0

        for key in ans_dict:
            device = ans_dict[key]

            t = d.add_table(rows=3, cols=4, style="Table Grid")
            t.cell(0, 1).merge(t.cell(0, 3))
            t.cell(1, 1).merge(t.cell(1, 3))
            t.cell(2, 1).merge(t.cell(2, 3))

            t.cell(0, 0).text = '设备名'
            t.cell(0, 1).text = device['设备名']
            t.cell(1, 0).text = 'SN'
            t.cell(1, 1).text = device['SN']
            t.cell(2, 0).text = '型号'
            t.cell(2, 1).text = device['型号']

            t = d.add_table(rows=3, cols=4, style="Table Grid")
            t.cell(0, 0).text = '角色'
            t.cell(0, 1).text = device['角色']
            t.cell(0, 2).text = '楼层'
            t.cell(0, 3).text = device['楼层']
            t.cell(1, 0).text = '机柜'
            t.cell(1, 1).text = device['机柜']
            t.cell(1, 2).text = 'U数'
            t.cell(1, 3).text = device['U数']
            t.cell(2, 0).text = '管理IP'
            t.cell(2, 1).text = device['管理IP']
            t.cell(2, 2).text = '管理接口'
            t.cell(2, 3).text = device['管理接口']

            d.add_paragraph('')
            sum_label += 1

            if sum_label >= labels_per_page:
                d.add_paragraph('')
                d.add_paragraph('')
                d.add_paragraph('')
                sum_label = 0

            d.save(word_path)

        logger.info(f'完成 {target} 的标签卡转换')
